import os
import time

import fitz  # PyMuPDF
import docx
import ollama

from backend.config import settings, logger


# ── File Parsing ─────────────────────────────────────────

def parse_file(file_path: str) -> str:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_size = os.path.getsize(file_path)
    max_size = settings.MAX_FILE_SIZE_MB * 1024 * 1024

    if file_size > max_size:
        raise ValueError(
            f"File size ({file_size / 1024 / 1024:.1f}MB) exceeds "
            f"maximum ({settings.MAX_FILE_SIZE_MB}MB)"
        )

    extension = os.path.splitext(file_path)[1].lower()
    logger.info(f"Parsing file: {file_path} (type: {extension}, size: {file_size} bytes)")

    parsers = {
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".txt": _parse_txt,
    }

    parser = parsers.get(extension)
    if parser is None:
        raise ValueError(
            f"Unsupported file type: {extension}. "
            f"Allowed: {', '.join(settings.ALLOWED_EXTENSIONS)}"
        )

    text = parser(file_path)
    logger.info(f"Parsed {len(text)} characters from {file_path}")
    return text


def _parse_pdf(file_path: str) -> str:
    try:
        document = fitz.open(file_path)
    except Exception as e:
        raise ValueError(f"Cannot open PDF: {str(e)}")

    pages = []
    for page_num in range(len(document)):
        page = document[page_num]
        text = page.get_text()
        if text.strip():
            pages.append(text)

    document.close()

    if not pages:
        raise ValueError(
            "PDF contains no extractable text. "
            "It may be scanned/image-only. OCR is not yet supported."
        )

    return "\n\n".join(pages)


def _parse_docx(file_path: str) -> str:
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        raise ValueError(f"Cannot open DOCX: {str(e)}")

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    if not paragraphs:
        raise ValueError("DOCX contains no extractable text.")

    return "\n\n".join(paragraphs)


def _parse_txt(file_path: str) -> str:
    encodings = ["utf-8", "latin-1", "cp1252"]

    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
            if content.strip():
                return content
        except UnicodeDecodeError:
            continue

    raise ValueError("Cannot read TXT file with any supported encoding.")


# ── Chunking ─────────────────────────────────────────────

def chunk_text(text: str) -> list[dict]:
    text = text.strip()
    if not text:
        raise ValueError("Cannot chunk empty text.")

    chunks = []
    start = 0

    while start < len(text):
        end = min(start + settings.CHUNK_SIZE, len(text))

        if end < len(text):
            adjusted_end = _find_sentence_break(text, start, end)
            if adjusted_end > start:
                end = adjusted_end

        chunk_content = text[start:end].strip()

        if chunk_content:
            chunks.append({
                "index": len(chunks),
                "text": chunk_content,
                "char_start": start,
                "char_end": end,
            })

        next_start = end - settings.CHUNK_OVERLAP
        start = next_start if next_start > start else end

    logger.info(f"Created {len(chunks)} chunks from {len(text)} characters")
    return chunks


def _find_sentence_break(text: str, start: int, end: int) -> int:
    search_start = start + int((end - start) * 0.7)
    search_region = text[search_start:end]

    for delimiter in ["\n\n", "\n", ". ", "? ", "! ", "; "]:
        last_break = search_region.rfind(delimiter)
        if last_break != -1:
            return search_start + last_break + len(delimiter)

    return end


# ── Embedding with Ollama ────────────────────────────────

def create_embeddings(chunks: list[dict]) -> list[list[float]]:
    logger.info(f"Creating embeddings for {len(chunks)} chunks using Ollama")

    embeddings = []

    for i, chunk in enumerate(chunks):
        for attempt in range(1, settings.MAX_RETRIES + 1):
            try:
                response = ollama.embed(
                    model=settings.EMBEDDING_MODEL,
                    input=chunk["text"],
                )
                embedding = response["embeddings"][0]
                embeddings.append(embedding)

                if (i + 1) % 5 == 0 or i == len(chunks) - 1:
                    logger.info(f"Embedded {i + 1}/{len(chunks)} chunks")

                break

            except Exception as e:
                logger.warning(f"Embedding failed for chunk {i} (attempt {attempt}): {e}")
                if attempt == settings.MAX_RETRIES:
                    raise RuntimeError(
                        f"Failed to embed chunk {i}. Is Ollama running? Run: ollama serve"
                    )
                time.sleep(2 ** attempt)

    logger.info(f"All {len(embeddings)} embeddings created successfully")
    return embeddings


# ── Full Pipeline ────────────────────────────────────────

def process_document(file_path: str) -> tuple[list[dict], list[list[float]], str]:
    logger.info(f"Starting document processing: {file_path}")

    full_text = parse_file(file_path)
    chunks = chunk_text(full_text)
    embeddings = create_embeddings(chunks)

    logger.info(
        f"Processing complete: {len(full_text)} chars → "
        f"{len(chunks)} chunks → {len(embeddings)} embeddings"
    )

    return chunks, embeddings, full_text